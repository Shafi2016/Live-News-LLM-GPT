import streamlit as st
import requests
import time
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI
import io
import logging
import sys
import re
import html
import yaml
from yaml.loader import SafeLoader
import streamlit_authenticator as stauth

# Set up logging
logging.basicConfig(level=logging.INFO, stream=sys.stdout,
                    format='%(asctime)s - %(levelname)s - %(message)s')

# Set page configuration (this should be at the top of your app)
st.set_page_config(layout="wide")  # Use wide layout for better readability

# Function to perform a Google search query using SerpAPI
def search_query(query, api_key):
    try:
        params = {
            "engine": "google",
            "q": query,
            "api_key": api_key
        }
        response = requests.get("https://serpapi.com/search", params=params)
        response.raise_for_status()  # Raise an error for bad responses
        return response.json()
    except requests.exceptions.RequestException as e:
        st.error(f"Error during search: {e}")
        logging.error(f"Error during search: {e}", exc_info=True)
        return None

# Function to extract relevant information from the search results
def extract_relevant_info(search_results):
    if search_results is None:
        return "", [], {}
    
    snippets = []
    links = []
    citations = {}
    
    for index, result in enumerate(search_results.get('organic_results', []), 1):
        snippet = result.get('snippet')
        link = result.get('link')
        
        if snippet and link:
            snippets.append(snippet)
            links.append(link)
            citations[f"[{index}]"] = link
    
    context = " ".join(snippets)  # Join all snippets with spaces to keep them readable together
    return context, links, citations

# Function to query GPT model with the provided question and context
def ask_gpt(question, context, client, citations, model_choice):
    try:
        # Adding the context as an explicit introduction in the answer
        introduction = f" \n{context}\n\n"
        full_question = f"{question}\n\nPlease use the following citation format when referencing sources: [1], [2], etc. The citations should correspond to the following references:\n"
        for citation, link in citations.items():
            full_question += f"{citation}: {link}\n"
        
        messages = [
            {"role": "system", "content": "You are a helpful assistant. When providing information, use the given citation format to reference sources."},
            {"role": "user", "content": full_question}
        ]
        
        response = client.chat.completions.create(
            model=model_choice,
            messages=messages,
            temperature=0,
            max_tokens=4000
        )
        
        answer = response.choices[0].message.content
        
        # Combine the introduction and the generated answer
        combined_answer = introduction + answer
        return combined_answer
    
    except Exception as e:
        st.error(f"Error during GPT query: {e}")
        logging.error(f"Error during GPT query: {e}", exc_info=True)
        return "An error occurred while processing the request."

# Function to add a hyperlink to a paragraph
def add_hyperlink(paragraph, url, text):
    # Ensure the URL is valid
    if not url.startswith(('http://', 'https://')):
        url = 'http://' + url  # Prepend http:// if no protocol is provided

    # This gets access to the document.xml.rels file and gets a new relationship id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    # Join all the xml elements together and add the required text to the w:r element
    new_run.append(r_pr)
    new_run.text = text
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink

# Function to save the response, context, and references to a DOCX file
def save_to_docx(answer, citations):
    doc = DocxDocument()
    
    # Add answer as a section (Introduction is already integrated in the answer)
    doc.add_heading("Answer", level=1)
    
    # Split the answer into paragraphs
    paragraphs = answer.split('\n')
    
    for paragraph in paragraphs:
        p = doc.add_paragraph()
        # Split the paragraph by citation markers
        parts = re.split(r'(\[\d+\])', paragraph)
        for part in parts:
            if part in citations:
                # This is a citation, add it as a hyperlink
                add_hyperlink(p, citations[part], part)
            else:
                # This is regular text, add it normally
                p.add_run(part)
    
    # Add references as a section
    doc.add_heading("References", level=1)
    for citation, link in citations.items():
        p = doc.add_paragraph(f"{citation} ")
        add_hyperlink(p, link, link)
    
    # Save the document to a BytesIO object
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file

# Function to clean and format the text
def clean_text(text):
    # Remove excessive asterisks and periods
    text = re.sub(r'\*+', '*', text)
    text = re.sub(r'\.+', '.', text)
    
    # Add spaces after periods if missing
    text = re.sub(r'\.(?=[A-Z])', '. ', text)
    
    # Decode HTML entities
    text = html.unescape(text)
    
    # Remove any remaining weird characters
    text = re.sub(r'[^\w\s.,;:!?*\[\]()"-]', '', text)
    
    return text

# Function to convert citations to clickable links in Markdown
def make_citations_clickable(text, citations):
    for citation, link in citations.items():
        text = text.replace(citation, f"[{citation}]({link})")
    return text

# Function to format the answer in nice Markdown
def format_answer_markdown(answer, citations):
    # Clean the text
    clean_answer = clean_text(answer)
    
    # Make citations clickable
    clickable_answer = make_citations_clickable(clean_answer, citations)
    
    # Remove the non-blue "Introduction" header if present
    clickable_answer = re.sub(r'\*\*Introduction\*\*\n*', '', clickable_answer)
    
    # Split the answer into introduction and main content
    parts = clickable_answer.split("\n\n", 1)
    introduction = parts[0] if len(parts) > 1 else ""
    main_content = parts[1] if len(parts) > 1 else parts[0]
    
    # Format the markdown
    formatted_answer = f"""
<h2 style='color: #0066cc;'>Introduction</h2>

 {introduction}

<h2 style='color: #0066cc;'>Analysis</h2>

{main_content}

"""
    return formatted_answer

# Load credentials and configuration from Streamlit Secrets
credentials = yaml.safe_load(st.secrets["general"]["credentials"])
cookie_name = st.secrets["general"]["cookie_name"]
cookie_key = st.secrets["general"]["cookie_key"]
cookie_expiry_days = st.secrets["general"]["cookie_expiry_days"]

# Create an authenticator object with hashed passwords
authenticator = stauth.Authenticate(
    credentials,
    cookie_name,
    cookie_key,
    cookie_expiry_days,
    None  # No preauthorized emails in this example
)

# Display the login form in the sidebar
name, authentication_status, username = authenticator.login('main')

if authentication_status:
    # Successful login
    authenticator.logout('Logout', 'main')  # Added logout option
    st.write(f'Welcome *{name}*')
    
    # Main app content goes here
 

    # Custom CSS to expand the main content area and reduce white space
    st.markdown("""
        <style>
            .main {
                max-width: 100%;
                margin: 0 auto;
                padding-left: 20px;
                padding-right: 20px;
            }
            .sidebar .element-container img {
                margin-left: auto;
                margin-right: auto;
                display: block;
            }
            .sidebar {
                width: 25%;
            }
        </style>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div style="background-color:#0077a8;padding:10px;border-radius:5px;width:100%;max-width:1200px;margin: 0 auto;">
            <h1 style="color:white;">SpotLight News</h1>
            <p style="color:white; font-size:17px; white-space: nowrap; text-align:left;">Spotlight the stories that matter with quick summaries and curated insights, delivered instantly.</p>
        </div>
        <div style="margin-bottom: 30px;"></div>
    """, unsafe_allow_html=True)

    # Streamlit app
    def main():
        try:
            st.title("")

            # Access API keys from Streamlit Secrets
            try:
                openai_api_key = st.secrets["general"]["OPENAI_API_KEY"]
                serpapi_api_key = st.secrets["general"]["SERPAPI_API_KEY"]
            except KeyError as e:
                st.error(f"Missing API key in Streamlit Secrets: {e}")
                logging.error(f"Missing API key in Streamlit Secrets: {e}")
                return

            # Initialize OpenAI client
            client = OpenAI(api_key=openai_api_key)

            # Sidebar for settings
            with st.sidebar:
                st.header("Settings")
                model_choice = st.selectbox("Select GPT Model", ["gpt-4o-mini", "gpt-4o"])

            # User input for search query (compulsory)
            query = st.text_input("Enter your search query (required):")

            # User input for analysis question (optional)
            question = st.text_area("Enter your analysis question (optional):")

            if st.button("Analyze"):
                if query:
                    with st.spinner("Searching and analyzing..."):
                        try:
                            # Perform the search query
                            search_results = search_query(query, serpapi_api_key)
                            
                            # Extract the context, relevant links, and citations from the search results
                            context, links, citations = extract_relevant_info(search_results)
                            
                            # Pause for a moment to ensure all context is retrieved
                            time.sleep(5)
                            
                            # If no specific question is provided, use a default one
                            if not question:
                                question = f"Provide a comprehensive summary and analysis of the information related to: {query}"
                            
                            # Get the answer from GPT using the context as the introduction
                            answer = ask_gpt(question, context, client, citations, model_choice)
                            
                            # Format the answer in nice Markdown
                            formatted_answer = format_answer_markdown(answer, citations)
                            
                            # Display the formatted answer
                            st.markdown(formatted_answer, unsafe_allow_html=True)
                            
                            # Display references
                            st.markdown("<h2 style='color: #0066cc;'>References</h2>", unsafe_allow_html=True)
                            for citation, link in citations.items():
                                st.markdown(f"{citation} [{link}]({link})")
                            
                            # Save the answer and references to a DOCX file
                            docx_file = save_to_docx(clean_text(answer), citations)
                            
                            # Provide a download button for the DOCX file
                            st.download_button(
                                label="Download Analysis as DOCX",
                                data=docx_file,
                                file_name="analysis_report.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        except Exception as e:
                            st.error(f"An error occurred during analysis: {e}")
                            logging.error(f"Error during analysis: {e}", exc_info=True)
                else:
                    st.warning("Please enter a search query.")
        except Exception as e:
            st.error(f"An unexpected error occurred: {e}")
            logging.error(f"Unexpected error: {e}", exc_info=True)

    if __name__ == "__main__":
        logging.info("Starting the Streamlit app")
        main()
        logging.info("Streamlit app execution completed")

elif authentication_status == False:
    st.error('Username/password is incorrect')

elif authentication_status is None:
    st.warning('Please enter your username and password')